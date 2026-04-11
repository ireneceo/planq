"""
텍스트 추출기.

입력: (content_type, body bytes, filename_hint)
출력: ExtractResult(text, title)

지원:
- HTML/XHTML → trafilatura (본문 + 제목)
- PDF → pdfplumber (텍스트 레이어만. 스캔본은 ExtractError)
- DOCX → python-docx
- TXT/MD → 직접 디코딩

블로킹 I/O(파서) 는 asyncio.to_thread 로 감싸서 호출자가 이벤트 루프를 블록하지 않도록 한다.
"""
import asyncio
import io
from dataclasses import dataclass
from typing import Optional


class ExtractError(Exception):
  """텍스트 추출 실패. message 는 사용자 노출 가능한 사유."""


@dataclass
class ExtractResult:
  text: str
  title: Optional[str]


# ─────────────────────────────────────────────────────────
# 확장자 / MIME → extractor 디스패치
# ─────────────────────────────────────────────────────────

EXT_MAP = {
  'html': 'html',
  'htm': 'html',
  'pdf': 'pdf',
  'docx': 'docx',
  'doc': 'docx',  # python-docx 는 .doc 실패할 수 있음 (그 경우 ExtractError)
  'txt': 'txt',
  'md': 'txt',
}

MIME_MAP = {
  'text/html': 'html',
  'application/xhtml+xml': 'html',
  'application/pdf': 'pdf',
  'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
  'application/msword': 'docx',
  'text/plain': 'txt',
  'text/markdown': 'txt',
}


def _resolve_kind(content_type: Optional[str], filename_hint: Optional[str]) -> str:
  if content_type:
    kind = MIME_MAP.get(content_type.split(';', 1)[0].strip().lower())
    if kind:
      return kind
  if filename_hint and '.' in filename_hint:
    ext = filename_hint.rsplit('.', 1)[1].lower()
    kind = EXT_MAP.get(ext)
    if kind:
      return kind
  raise ExtractError(f'지원하지 않는 형식 (content_type={content_type}, filename={filename_hint})')


# ─────────────────────────────────────────────────────────
# 개별 추출기 (모두 동기 → to_thread 로 감쌈)
# ─────────────────────────────────────────────────────────

def _extract_html_sync(body: bytes) -> ExtractResult:
  import trafilatura
  try:
    html = body.decode('utf-8', errors='replace')
  except Exception:
    raise ExtractError('HTML 디코딩 실패')
  extracted = trafilatura.extract(
    html,
    include_comments=False,
    include_tables=True,
    favor_precision=True,
    deduplicate=True,
  )
  if not extracted or not extracted.strip():
    raise ExtractError('본문을 추출할 수 없는 HTML 입니다')
  # 제목은 metadata 에서
  title = None
  try:
    meta = trafilatura.extract_metadata(html)
    if meta and getattr(meta, 'title', None):
      title = meta.title.strip()[:200]
  except Exception:
    pass
  return ExtractResult(text=extracted.strip(), title=title)


def _extract_pdf_sync(body: bytes) -> ExtractResult:
  import pdfplumber
  texts = []
  title = None
  try:
    with pdfplumber.open(io.BytesIO(body)) as pdf:
      meta = pdf.metadata or {}
      raw_title = meta.get('Title') or meta.get('title')
      if raw_title:
        title = str(raw_title).strip()[:200] or None
      for page in pdf.pages:
        try:
          t = page.extract_text()
        except Exception:
          t = None
        if t:
          texts.append(t)
  except Exception as e:
    raise ExtractError(f'PDF 파싱 실패: {type(e).__name__}')

  combined = '\n\n'.join(t.strip() for t in texts if t and t.strip())
  if not combined:
    raise ExtractError('PDF에서 텍스트를 추출할 수 없습니다 (스캔 PDF는 아직 지원하지 않습니다)')
  return ExtractResult(text=combined, title=title)


def _extract_docx_sync(body: bytes) -> ExtractResult:
  try:
    from docx import Document
  except ImportError:
    raise ExtractError('DOCX 파서가 설치되지 않았습니다')
  try:
    doc = Document(io.BytesIO(body))
  except Exception as e:
    raise ExtractError(f'DOCX 파싱 실패: {type(e).__name__}')

  paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
  # 표 본문도 수집
  for table in doc.tables:
    for row in table.rows:
      row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text and cell.text.strip())
      if row_text:
        paragraphs.append(row_text)

  if not paragraphs:
    raise ExtractError('DOCX에서 텍스트를 추출할 수 없습니다')

  text = '\n\n'.join(paragraphs)
  title = None
  try:
    core = doc.core_properties
    if core and core.title:
      title = core.title.strip()[:200] or None
  except Exception:
    pass
  return ExtractResult(text=text, title=title)


def _extract_txt_sync(body: bytes) -> ExtractResult:
  # UTF-8 우선, 실패 시 cp949 (한국어 로컬) → latin-1 fallback
  for enc in ('utf-8', 'utf-8-sig', 'cp949', 'euc-kr', 'latin-1'):
    try:
      text = body.decode(enc)
      break
    except UnicodeDecodeError:
      continue
  else:
    raise ExtractError('텍스트 디코딩 실패')
  text = text.strip()
  if not text:
    raise ExtractError('빈 텍스트 파일입니다')
  return ExtractResult(text=text, title=None)


# ─────────────────────────────────────────────────────────
# Public entry
# ─────────────────────────────────────────────────────────

async def extract(body: bytes, content_type: Optional[str] = None, filename_hint: Optional[str] = None) -> ExtractResult:
  kind = _resolve_kind(content_type, filename_hint)
  if kind == 'html':
    return await asyncio.to_thread(_extract_html_sync, body)
  if kind == 'pdf':
    return await asyncio.to_thread(_extract_pdf_sync, body)
  if kind == 'docx':
    return await asyncio.to_thread(_extract_docx_sync, body)
  if kind == 'txt':
    return await asyncio.to_thread(_extract_txt_sync, body)
  raise ExtractError(f'지원하지 않는 형식: {kind}')
